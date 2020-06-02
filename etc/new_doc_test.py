#该脚本主要用于doc文档利用百度搜索引擎自动查寻并返回类似文本进行高亮批注
#只能读取[.docx]文件，不能读取[.doc]文件
from docx import Document
#此处在r'xxxx'中输入绝对路径
path = r'C:\Users\admin\Desktop\未完整版：美腾科技所处行业分析报告（菁亿投顾-2020-05-25).docx'

def get_docx_paragraph(path):
    paragraph = []
    document = Document(path)
    for i in document.paragraphs:
        paragraph.append(i.text)    #有陷阱,空list是nonetype,nonetype不能append所以不能用赋值的方法
    paragraph = [i for i in paragraph if i != '']   #列表生成式去空集比较简单
    return paragraph

#重新输出文本
paragraph = get_docx_paragraph(path)
new_doc = Document()
for i in paragraph:
    new_doc.add_paragraph(i)
new_doc.save(r'C:\Users\admin\Desktop\test.docx')